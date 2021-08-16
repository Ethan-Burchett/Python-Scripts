
#fileSelectFile

#from docx import Document
#from docx.shared import Inches

import subprocess


from posixpath import split
import tkinter as tk
from tkinter import simpledialog, filedialog
from tkinter.constants import TRUE
#from typing_extensions import ParamSpec

from docx.shared import Inches  

import os

import sys

from docxfunctions import *

from docx.shared import Pt

from docx import Document
from docx.shared import Pt


root = tk.Tk()

root.withdraw() # this hides the root window


#string constants
sentence_fragment_to_keep = 'For this period, our involvement was limited to'
app_title = "Weekly Report Helper"
last_line = "Respectfully Submitted,"


date_range = "July 4 through 10, 2021"
month_date_year = "July 13, 2021"
#date_range = simpledialog.askstring(title = app_title,prompt="What is the date range?\nEX: April 4 through 10, 2021")

#month_date_year = simpledialog.askstring(title = app_title,prompt="What is the date?\nEX: April 13, 2021")



# Build a list of tuples for each file type the file dialog should display
my_filetypes = [('all files', '.*'), ('text files', '.txt')]

temp = 0

while temp != -1:

    increment_report_num = 0
    temp = input("continue?")

    file_path = filedialog.askopenfile(parent=root,initialdir=os.getcwd(),title="Please select a file:", filetypes=my_filetypes)  # returns fancy object that has the path as the name. 

    if file_path.name.endswith('.doc'):
        print(".doc")
        doc_to_docx(file_path.name)
        path = file_path.name + "x"
    else:
        path = file_path.name


    if file_path is not None:
        print(file_path.name)


    #date_range = ""
    #month_date_year = ""
    #date_range = simpledialog.askstring(title = app_title,prompt="What is the date range?\nEX: April 4 through 10, 2021")

    #month_date_year = simpledialog.askstring(title = app_title,prompt="What is the date?\nEX: April 13, 2021")

    document = Document(path) # opens the document spec. by the path 

    lineNumber = 0
    report_num = 0

    #styles = document.styles['Normal']
    #font = styles.font
    #font.size = Pt(11)

    #style = document.styles['Normal']
    #font = style.font
    #font.size = Pt(11)
    



    for para in document.paragraphs:
        #print(para.text + " " + str(lineNumber))
        lineNumber = lineNumber + 1
        
        line = para.text

        if "\t\tRepo" in line:
            print("found report line" + line)
            split_line = line.split("#")
            #print(split_line)
            report_num = split_line[1]
            #print(report_num)
            increment_report_num = int(report_num) + 1
            #print(increment_report_num)
            text = "\t\t\tReport#" + str(increment_report_num)
            para.text = ""
            #para.style = 'Normal'
            run = para.add_run(text).bold = True
            #run.style = 'Normal'
            #run.font.size = Pt(12)
            
            #para.add_run(text).font.size = Pt(11)
        # para.font.size = Pt(11)
            #para.text.run.font = Pt(11)
            #font = para.font
            #font.size = Pt(11)
            #para.run.size = Pt(11)
            #para.font.size = Pt(11)


        if lineNumber == 2:
            print("at date-month? ")
            #para.text = para.text + "****"
            line = para.text 
            #print(line)
            split_line = line.split("\t")
            #print(split_line[1])
            #split_line[1] = month_date_year
            line = split_line[0] + "\t" + month_date_year
            #print(line)
            para.text = ""
            para.add_run(line).bold = False

        if "through" in line and ", 2021" in line or "through" in line and ", 202" in line:
            print("at date range")
            print(line)
            para.text = ""
            line = "\t\t\t" + date_range
            para.add_run(line).bold = True
            
        #if string_to_search.find(sentence_fragment_to_keep) == 1:
        #   print("******found it********")

        if sentence_fragment_to_keep in para.text:
            print("******found it********")   # get index after this point 
            para.text = sentence_fragment_to_keep
            #para.add_run(sentence_fragment_to_keep)

    #print(sentence_fragment_to_keep)

    #print(date_range)
    #print(month_date_year)

    #for para in document.paragraphs:
    #   para.style = document.style['Normal']

    #for run in para.runs:
    #   run.font.size = Pt(11)


    file_name = file_path.name
    split_file_name = file_name.split("#")

    #name = split_file_name[0] + "#" + str(int(split_file_name[1]) + 1)

    
    name = split_file_name[0] + "#" + str(increment_report_num) + ".docx"
    print(name)
    document.save(name)